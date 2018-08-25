"""
	将doc文件转化成txt文件作为语料库使用
"""
import os
import docx
from docx import Document


"""关键词生成一个txt文件"""

"""

path = './Datasets/key_words/'
file_list = os.listdir(path)
f = open("./courpus/dict/dict.txt", 'a+')

for file in file_list:
	file_path = path + file
	document = Document(file_path)
	for paragraph in document.paragraphs:
		#print(paragraph.text)
		for _ in paragraph.text.split():
			f.write(_)
			f.write('\n')

f.close()
"""

"""将数据文档转化成txt文件"""

"""

path = './Datasets/ch_data/'
file_list = os.listdir(path)
f = open("./corpus/1/data.txt", 'a+')
for file in file_list:
	file_path = path + file
	document = Document(file_path)
	for paragraph in document.paragraphs:
		f.write(paragraph.text)
f.close()
"""

"""将关键词按年分类转化成txt"""

path = './Datasets/key_words/'
file_list = os.listdir(path)
for file in file_list:
	f = open('./Datasets/key_words_txt/' + file[0:4] + '.txt', 'w')
	file_path = path + file
	document = Document(file_path)
	for paragraph in document.paragraphs:
		f.write(paragraph.text)
		f.write('\n')
	f.close()
